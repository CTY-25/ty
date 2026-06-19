"""
生成共享单车数据分析和Streamlit应用完整程序代码文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

BASE_DIR = r'C:\Users\32404\Desktop\新建文件夹'

doc = Document()

# 设置正常样式
style = doc.styles['Normal']
font = style.font
font.name = 'Consolas'
font.size = Pt(10)
style.paragraph_format.line_spacing = 1.2

# 标题样式
h_style = doc.styles['Heading 1']
h_font = h_style.font
h_font.name = '黑体'
h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h_font.size = Pt(16)
h_font.bold = True

h2_style = doc.styles['Heading 2']
h2_font = h2_style.font
h2_font.name = '黑体'
h2_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h2_font.size = Pt(13)
h2_font.bold = True

def add_code_block(code_text, title=None):
    """添加代码块"""
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run('# ======== ' + title + ' ========')
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.bold = True
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

def add_comment(comment_text):
    """添加注释性说明"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(comment_text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

def load_file(filename):
    """读取代码文件内容"""
    path = os.path.join(BASE_DIR, filename)
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

# ============================================================
# 标题
# ============================================================
doc.add_heading('附件1：程序代码', level=0)
add_comment('本文件包含共享单车使用模式与城市出行特征分析的全部Python源代码，包括数据预处理、分析、建模及Streamlit应用。')
add_comment('开发环境：Python 3.10+')
add_comment('依赖：streamlit, pandas, numpy, matplotlib, seaborn, plotly, scikit-learn, geopandas')
add_comment('')

# ============================================================
# 1. requirements.txt
# ============================================================
doc.add_heading('1. 项目依赖文件（requirements.txt）', level=1)
add_code_block(load_file('requirements.txt'), '项目依赖包列表')

# ============================================================
# 2. data_preprocessing.py
# ============================================================
doc.add_heading('2. 数据预处理模块（data_preprocessing.py）', level=1)
add_code_block(load_file('data_preprocessing.py'), '数据预处理模块完整代码')

# ============================================================
# 3. analysis_visualization.py
# ============================================================
doc.add_heading('3. 分析与可视化模块（analysis_visualization.py）', level=1)
add_code_block(load_file('analysis_visualization.py'), '数据分析与静态可视化模块')

# ============================================================
# 4. app.py
# ============================================================
doc.add_heading('4. Streamlit应用主文件（app.py）', level=1)
add_comment('以下是完整的Streamlit Web应用代码，整合了所有功能模块。')
add_code_block(load_file('app.py'), 'Streamlit主应用完整代码')

# ============================================================
# 5. utils.py
# ============================================================
doc.add_heading('5. 工具函数模块（utils.py）', level=1)
add_code_block(load_file('utils.py'), '工具函数模块')

# 保存
output_path = os.path.join(BASE_DIR, '附件1-程序代码.docx')
doc.save(output_path)
print(f'程序代码文档已保存至：{output_path}')
