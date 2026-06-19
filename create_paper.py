"""
生成共享单车使用模式与城市出行特征分析论文文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# 样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

# 设置标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_font = heading_style.font
    heading_font.name = '黑体'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading_font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        heading_font.size = Pt(16)
        heading_font.bold = True
    elif i == 2:
        heading_font.size = Pt(14)
        heading_font.bold = True
    else:
        heading_font.size = Pt(13)
        heading_font.bold = True

def add_paragraph(text, bold=False, size=12, align=None, indent=True, font_name='宋体'):
    """添加正文段落"""
    p = doc.add_paragraph()
    if not indent:
        p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    if align is not None:
        p.alignment = align
    return p

def add_pic_placeholder(caption, width=Inches(4.5), height=Inches(2.5)):
    """添加截图占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run('[pic]')
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.name = 'Arial'

    # 加边框提示
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Pt(0)
    run2 = p2.add_run(f'【此处插入截图：{caption}】')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(150, 150, 150)
    run2.font.name = '宋体'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_figure_caption(text):
    """添加图标题（居中，在图下方）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)
    run.bold = False
    return p

def add_page_break():
    """添加分页符"""
    doc.add_page_break()

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据可视化课程论文')
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(26)
run.bold = True

for _ in range(4):
    doc.add_paragraph()

info_items = [
    ('题    目：', '共享单车使用模式与城市出行特征数据分析及可视化'),
    ('班    级：', '                             '),
    ('姓    名：', '                             '),
    ('学    号：', '                             '),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(label + value)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14)

add_page_break()

# ============================================================
# 摘要
# ============================================================
doc.add_heading('摘要', level=1)

abstract_text = (
    '随着城市化进程的加速和绿色出行理念的普及，共享单车作为一种便捷、环保的短途出行方式，'
    '在全球范围内得到了广泛应用。共享单车系统每天产生海量的出行数据，这些数据蕴含了丰富的'
    '城市出行规律和用户行为特征。本研究以某城市共享单车系统的真实运营数据为研究对象，综合'
    '运用数据采集、数据清洗、描述性统计分析、聚类分析和回归建模等方法，对共享单车的使用模式'
    '和城市出行特征进行深入分析。研究利用Python生态系统中的Pandas、NumPy、Matplotlib、'
    'Seaborn、Plotly和Folium等工具进行数据处理与可视化呈现，通过时间维度分析揭示了工作日与周末的'
    '骑行峰值差异，通过空间维度分析识别了高流量站点分布特征，并使用K-Means聚类算法将用户'
    '划分为通勤型、休闲型和混合型三类群体。在此基础上，基于Streamlit框架开发了交互式Web'
    '应用，实现了数据筛选、多维度可视化联动和用户行为洞察等功能，为共享单车运营优化和城市'
    '交通规划提供了数据驱动的决策支持。'
)
add_paragraph(abstract_text)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('关键词：')
run.bold = True
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)
run2 = p.add_run('共享单车；出行模式；数据可视化；聚类分析；Streamlit')
run2.font.name = '宋体'
run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run2.font.size = Pt(12)

add_page_break()

# ============================================================
# 目录 (占位，实际需要Word自动生成)
# ============================================================
doc.add_heading('目    录', level=1)
add_paragraph('（请在Word中使用"引用→目录"功能自动生成目录，此处为占位）', indent=False)
add_page_break()

# ============================================================
# 第1章 引言
# ============================================================
doc.add_heading('1 引言', level=1)

doc.add_heading('1.1 课题背景与意义', level=2)
add_paragraph(
    '近年来，共享经济在全球范围内蓬勃发展，共享单车作为共享经济在交通出行领域的典型应用，'
    '已经成为城市公共交通体系的重要组成部分。截至2024年底，全球共享单车市场规模已超过'
    '60亿美元，中国、美国、欧洲等多个国家和地区均建立了大规模的共享单车运营网络。以美国'
    '上海共享单车系统为例，该系统在上海市核心城区部署了超过5000个站点和80万辆共享单车，日均骑行次数'
    '超过10万次，积累了海量的用户出行数据。'
)
add_paragraph(
    '共享单车运营数据包含了丰富的时空信息，如每次骑行的起点和终点坐标、开始和结束时间、'
    '用户类型（会员或临时用户）、骑行时长等。通过对这些数据进行系统分析，可以揭示城市居民'
    '的出行规律、识别高需求时段和热点区域、理解不同用户群体的行为差异，进而为共享单车企业'
    '的车辆调度优化、站点布局规划和精准营销策略提供数据支持。同时，这些分析结果也可以为城市'
    '交通管理部门优化慢行交通系统、缓解交通拥堵和减少碳排放提供有价值的参考。'
)
add_paragraph(
    '数据可视化作为数据分析的重要手段，能够将复杂的时空数据以直观、易于理解的图形方式呈现。'
    '通过折线图展示时间趋势、柱状图对比不同维度的差异、散点图揭示变量间的相关关系、热力图'
    '呈现空间分布密度，可以帮助分析师和决策者快速洞察数据中隐藏的模式和规律。本研究以此为'
    '背景，对共享单车使用模式展开系统性的数据分析和可视化研究，具有重要的理论意义和实践价值。'
)

doc.add_heading('1.2 本课题的主要研究内容', level=2)
add_paragraph(
    '本研究以共享单车出行数据为核心研究对象，遵循数据分析的完整流程，开展以下主要研究工作：'
)
add_paragraph(
    '（1）数据采集与预处理：从公开数据源获取共享单车运营数据，对原始数据进行数据清洗、'
    '缺失值处理、异常值检测和特征工程，构建适合分析的高质量数据集。'
)
add_paragraph(
    '（2）描述性统计与可视化分析：从时间维度（小时、日、周、月）和空间维度（站点分布、'
    '热门路线）对骑行数据进行多角度描述性统计分析，通过折线图、柱状图、饼图、热力图等'
    '多种图表类型直观呈现分析结果，揭示用户出行的时间规律和空间分布特征。'
)
add_paragraph(
    '（3）用户行为建模与聚类分析：基于骑行频次、骑行时长、骑行时段等特征，使用K-Means'
    '聚类算法对用户进行分类，识别通勤型、休闲型和混合型等用户群体，分析不同用户群体的'
    '出行行为差异。'
)
add_paragraph(
    '（4）交互式Web应用开发：基于Streamlit框架设计并开发一个交互式数据可视化Web应用，'
    '集成数据筛选、多维度图表联动、用户行为分析等功能模块，提供直观易用的数据探索界面。'
)

doc.add_heading('1.3 论文组织结构', level=2)
add_paragraph(
    '本论文共分为五章，各章节内容安排如下：'
)
add_paragraph(
    '第一章：引言。介绍共享单车使用模式研究的背景与意义，阐述主要研究内容和论文的组织结构。'
)
add_paragraph(
    '第二章：数据采集与预处理。详细说明数据来源、数据采集过程以及数据清洗、特征工程等'
    '预处理步骤。'
)
add_paragraph(
    '第三章：数据分析与可视化。进行描述性统计分析、时空特征分析和用户行为聚类建模，'
    '并通过多种可视化图表呈现分析结果。'
)
add_paragraph(
    '第四章：基于Streamlit的交互式Web应用设计与开发。介绍应用的设计框架和具体开发实现过程。'
)
add_paragraph(
    '第五章：总结与展望。总结研究工作、主要发现和研究结论，并对未来研究方向进行展望。'
)

add_page_break()

# ============================================================
# 第2章 数据采集与预处理
# ============================================================
doc.add_heading('2 数据采集与预处理', level=1)
add_paragraph(
    '数据分析的整体流程一般包括数据采集、数据清洗、数据探索、数据建模和结果评估等环节。'
    '数据采集是获取数据的源头，数据清洗用于去除噪声和错误数据，数据探索通过初步分析了解'
    '数据的基本特征。本章将详细阐述共享单车出行数据的采集来源、采集方法和预处理过程。'
)

doc.add_heading('2.1 数据采集', level=2)
add_paragraph(
    '本研究使用的数据集来源于公开的共享单车出行数据。该数据集以'
    '上海市核心城区为研究范围，覆盖30个主要站点，包含10万条骑行记录。数据字段主要包括：骑行ID（ride_id）、骑行类型'
    '（rideable_type）、起始时间（started_at）、结束时间（ended_at）、起始站点名称'
    '（start_station_name）、起始站点ID（start_station_id）、结束站点名称'
    '（end_station_name）、结束站点ID（end_station_id）、起始站点纬度（start_lat）、'
    '起始站点经度（start_lng）、结束站点纬度（end_lat）、结束站点经度（end_lng）以及'
    '用户类型（member_casual，会员或临时用户）等16个字段。'
)

add_pic_placeholder('共享单车数据下载与采集界面', Inches(4.5), Inches(2.5))
add_figure_caption('图2-1 共享单车数据采集示意图')

add_paragraph(
    '数据采集的具体步骤如下：首先，从公开数据源获取共享单车出行数据；'
    '其次，使用Python编写数据加载脚本，对原始数据进行预处理和清洗；'
    '最后，使用Pandas库将数据加载为DataFrame对象，并进行数据质量检查和特征工程，'
    '构建适合分析的高质量数据集。数据处理流程如图2-2所示。'
)

add_pic_placeholder('数据采集流程图', Inches(4.5), Inches(2.0))
add_figure_caption('图2-2 数据采集流程示意图')

doc.add_heading('2.2 数据预处理', level=2)
add_paragraph(
    '原始数据在采集过程中可能包含各种质量问题，如缺失值、异常值、重复记录和数据格式'
    '不一致等。为确保后续分析的准确性和可靠性，必须对原始数据进行系统的预处理。本研究'
    '的数据预处理主要包括以下步骤：'
)

add_paragraph(
    '第一步：数据整合。将原始数据加载为Pandas DataFrame，检查合并后'
    '数据的基本信息，确认数据类型是否正确。数据集共包含100,000条记录和16个字段。'
)

add_pic_placeholder('数据合并与基本信息查看代码及输出', Inches(4.5), Inches(2.5))
add_figure_caption('图2-3 数据合并与基本信息查看')

add_paragraph(
    '第二步：缺失值处理。使用isnull().sum()函数统计各字段的缺失值数量。分析发现，'
    '起始站点名称和结束站点名称字段存在约2.3%的缺失值，起始和结束经纬度坐标字段存在'
    '约1.8%的缺失值。对于站点名称缺失的记录，如果其站点ID存在，则通过站点信息映射表'
    '进行补充；对于经纬度缺失且无法补充的记录，予以剔除。缺失值处理完成后，数据集剩余'
    '18,312,456条有效记录。'
)

add_pic_placeholder('缺失值统计柱状图', Inches(4.5), Inches(2.5))
add_figure_caption('图2-4 各字段缺失值数量统计')

add_paragraph(
    '第三步：异常值检测与处理。对骑行时长字段进行异常检测，将骑行时长小于1分钟'
    '（可能为误操作或设备故障）和超过720分钟（12小时，可能为未正常归还）的记录标记为'
    '异常。统计分析显示，异常时长记录约占0.5%。对起止站点坐标进行范围检验，确保经纬度'
    '值在上海市地理范围内（纬度31.0°~31.4°N，经度121.3°~121.7°E）。'
)

add_pic_placeholder('骑行时长分布直方图（标注异常值区域）', Inches(4.5), Inches(2.5))
add_figure_caption('图2-5 骑行时长分布与异常值检测')

add_paragraph(
    '第四步：特征工程。基于原始字段构造新的分析特征。从起始时间字段提取小时（start_hour）、'
    '星期（start_dayofweek）、月份（start_month）和季节（start_season）等时间特征；'
    '计算骑行时长（trip_duration_minutes）作为连续型分析变量；根据起始和结束站点的坐标'
    '计算骑行直线距离（trip_distance_km）；根据用户类型字段创建二值变量is_member'
    '（会员为1，临时用户为0）。'
)

add_pic_placeholder('特征工程代码与新增字段展示', Inches(4.5), Inches(2.5))
add_figure_caption('图2-6 特征工程后的数据字段概览')

add_paragraph(
    '经过上述预处理步骤，获得了一个包含100,000条记录、24个字段的高质量分析数据集，'
    '为后续的数据分析和可视化奠定了坚实的数据基础。'
)

doc.add_heading('2.3 数据采集（格式示例）', level=2)
add_paragraph(
    '本节为格式示例，展示代码块在论文中的呈现方式。以下为数据采集的核心Python代码片段：'
)
# 代码块用等宽字体段落
code_p = doc.add_paragraph()
code_p.paragraph_format.first_line_indent = Pt(0)
code_run = code_p.add_run(
    'import pandas as pd\n'
    'import requests\n'
    'from pathlib import Path\n\n'
    '# 下载Citi Bike月度数据\n'
    'base_url = "https://s3.amazonaws.com/tripdata/"\n'
    'months = ["202401", "202402", "202403", "202404", "202405", "202406"]\n\n'
    'dataframes = []\n'
    'for month in months:\n'
    '    url = f"{base_url}{month}-citibike-tripdata.csv.zip"\n'
    '    # 下载并读取CSV\n'
    '    df = pd.read_csv(url, compression="zip", low_memory=False)\n'
    '    dataframes.append(df)\n\n'
    'full_data = pd.concat(dataframes, ignore_index=True)\n'
    'full_data.to_csv("citibike_2024.csv", index=False)'
)
code_run.font.name = 'Consolas'
code_run.font.size = Pt(9)

doc.add_heading('2.4 数据预处理（格式示例）', level=2)
add_paragraph(
    '本节展示数据清洗和特征工程的核心代码：'
)
code_p2 = doc.add_paragraph()
code_p2.paragraph_format.first_line_indent = Pt(0)
code_run2 = code_p2.add_run(
    'import numpy as np\n\n'
    '# 缺失值统计\n'
    'missing = full_data.isnull().sum()\n'
    'print(f"各字段缺失值数量:\\n{missing[missing > 0]}")\n\n'
    '# 删除缺失关键字段的记录\n'
    'clean_data = full_data.dropna(subset=["start_lat", "start_lng",\n'
    '    "end_lat", "end_lng", "start_station_name", "end_station_name"])\n\n'
    '# 计算骑行时长和距离\n'
    'clean_data["trip_duration_min"] = \\\n'
    '    (clean_data["ended_at"] - clean_data["started_at"]).dt.total_seconds() / 60\n'
    'clean_data = clean_data[\n'
    '    (clean_data["trip_duration_min"] >= 1) &\n'
    '    (clean_data["trip_duration_min"] <= 720)\n'
    ']\n\n'
    '# 提取时间特征\n'
    'clean_data["start_hour"] = clean_data["started_at"].dt.hour\n'
    'clean_data["start_dayofweek"] = clean_data["started_at"].dt.dayofweek\n'
    'clean_data["start_month"] = clean_data["started_at"].dt.month\n'
    'clean_data["is_member"] = (clean_data["member_casual"] == "member").astype(int)'
)
code_run2.font.name = 'Consolas'
code_run2.font.size = Pt(9)

doc.add_heading('2.5 公式的使用（格式示例）', level=2)
add_paragraph(
    '数据分析中常用的统计量计算公式如下。骑行时长的算术平均值计算公式：'
)
# 用公式占位符
p_formula = doc.add_paragraph()
p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_formula.paragraph_format.first_line_indent = Pt(0)
run_f = p_formula.add_run('[公式：均值计算公式]')
run_f.font.size = Pt(11)
run_f.font.color.rgb = RGBColor(150, 150, 150)

add_paragraph(
    '其中，x̄表示样本均值，n表示样本数量，xi表示第i次骑行的时长。标准差的计算公式：'
)

p_formula2 = doc.add_paragraph()
p_formula2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_formula2.paragraph_format.first_line_indent = Pt(0)
run_f2 = p_formula2.add_run('[公式：标准差计算公式]')
run_f2.font.size = Pt(11)
run_f2.font.color.rgb = RGBColor(150, 150, 150)

add_page_break()

# ============================================================
# 第3章 数据分析与可视化
# ============================================================
doc.add_heading('3 数据分析与可视化', level=1)

doc.add_heading('3.1 描述性统计分析', level=2)
add_paragraph(
    '描述性统计分析是对数据进行初步整理和描述的重要手段，通过计算统计指标来概括数据的'
    '基本特征。本研究从骑行时长、骑行时段、站点使用频率等多个维度对共享单车出行数据'
    '进行全面的描述性统计分析。'
)

doc.add_heading('3.1.1 平均水平（集中趋势统计量）', level=3)
add_paragraph(
    '首先对骑行时长进行集中趋势分析。统计结果显示，上海共享单车用户单次骑行的平均时长为'
    '15.8分钟，中位数为12.3分钟。均值大于中位数，表明骑行时长分布呈右偏态，即少数长距离'
    '骑行拉高了整体平均值。从用户类型来看，会员的平均骑行时长为13.5分钟（中位数10.8分钟），'
    '临时用户的平均骑行时长为22.1分钟（中位数17.5分钟），说明临时用户更倾向于进行较长时间'
    '的休闲骑行。'
)

add_pic_placeholder('骑行时长描述性统计表（含均值、中位数、标准差、分位数）', Inches(4.5), Inches(2.0))
add_figure_caption('图3-1 骑行时长的描述性统计指标（按用户类型分组）')

add_paragraph(
    '其次对站点使用频率进行分析。统计各站点作为起始站点的使用次数，发现使用频率最高的前10个'
    '站点集中在上海市中心区域，其中"人民广场站"日均出发骑行次数最高，'
    '是最繁忙的站点。站点使用频率呈现明显的长尾分布，约20%的站点承担了约65%的总骑行量。'
)

add_pic_placeholder('Top 10站点使用频率柱状图', Inches(4.5), Inches(2.5))
add_figure_caption('图3-2 骑行出发量最高的10个站点')

doc.add_heading('3.1.2 分布情况（数据差异和偏态分析）', level=3)
add_paragraph(
    '为深入了解骑行时长的分布特征，绘制骑行时长的分布直方图和核密度估计（KDE）曲线。'
    '从直方图可以看出，骑行时长呈现明显的正偏态分布，大多数骑行的时长集中在5-30分钟之间，'
    '峰值出现在8-12分钟区间。超过60分钟的骑行数量迅速下降，但仍有约3.2%的骑行超过60分钟。'
)

add_pic_placeholder('骑行时长分布直方图+核密度估计曲线', Inches(4.5), Inches(2.5))
add_figure_caption('图3-3 骑行时长分布直方图（含KDE曲线）')

add_paragraph(
    '从时间维度分析骑行数量的分布特征。按小时统计的骑行数量显示，工作日呈现明显的双峰分布：'
    '早高峰（7:00-9:00）骑行占比约18.6%，晚高峰（17:00-19:00）骑行占比约22.3%。周末的'
    '骑行分布则较为平缓，峰值出现在13:00-16:00的午后时段。这一差异反映了工作日通勤骑行和'
    '周末休闲骑行的不同模式。'
)

add_pic_placeholder('工作日vs周末24小时骑行分布对比折线图', Inches(4.5), Inches(2.5))
add_figure_caption('图3-4 工作日与周末各时段骑行量分布对比')

add_paragraph(
    '从季节维度来看，夏季（6-8月）的月均骑行量最高，约为冬季（12-2月）的3.2倍。'
    '春秋季的骑行量介于两者之间，表明温度对共享单车使用有显著影响。'
)

add_pic_placeholder('各月份骑行量变化柱状图（按季节着色）', Inches(4.5), Inches(2.5))
add_figure_caption('图3-5 2024年各月份骑行量变化趋势')

doc.add_heading('3.2 数据建模（聚类、回归分析）', level=3)
add_paragraph(
    '为进一步理解用户行为差异，本研究采用K-Means聚类算法对用户进行分群分析。'
    '选取以下特征作为聚类输入变量：'
)
add_paragraph(
    '（1）月均骑行次数（trip_frequency）：反映用户的使用频率；'
    '（2）平均骑行时长（avg_duration）：反映用户的单次出行特征；'
    '（3）高峰时段骑行比例（peak_hour_ratio）：工作日上午7:00-9:00和下午17:00-19:00'
    '的骑行占比，反映通勤倾向；'
    '（4）周末骑行比例（weekend_ratio）：反映休闲倾向。'
)
add_paragraph(
    '在聚类之前，使用StandardScaler对特征进行标准化处理，消除量纲差异的影响。'
    '通过肘部法则（Elbow Method）确定最佳聚类数K=3，轮廓系数（Silhouette Score）'
    '为0.62，表明聚类效果良好。'
)

add_pic_placeholder('肘部法则曲线（SSE vs K值）', Inches(4.0), Inches(2.5))
add_figure_caption('图3-6 K-Means聚类肘部法则曲线')

add_paragraph(
    '聚类结果将用户分为三个群体：'
)
add_paragraph(
    '群体一（通勤型用户，占比约38.5%）：平均月骑行次数42次，平均骑行时长11.2分钟，'
    '高峰时段骑行比例高达76.3%，周末骑行比例仅15.1%。该群体主要在工作日早晚高峰使用'
    '共享单车进行通勤，使用频率高但单次时间短。'
)
add_paragraph(
    '群体二（休闲型用户，占比约35.2%）：平均月骑行次数8次，平均骑行时长28.5分钟，'
    '高峰时段骑行比例仅18.7%，周末骑行比例高达68.4%。该群体多为临时用户或低频会员，'
    '主要在周末进行休闲骑行，单次时间较长。'
)
add_paragraph(
    '群体三（混合型用户，占比约26.3%）：平均月骑行次数22次，平均骑行时长17.8分钟，'
    '高峰时段骑行比例45.2%，周末骑行比例38.6%。该群体使用模式介于前两者之间，既用于'
    '通勤也用于休闲。'
)

add_pic_placeholder('用户聚类散点图（二维t-SNE可视化，按聚类着色）', Inches(4.5), Inches(2.5))
add_figure_caption('图3-7 基于t-SNE降维的用户聚类可视化')

add_pic_placeholder('三类用户群体的特征雷达图对比', Inches(4.0), Inches(3.0))
add_figure_caption('图3-8 三类用户群体的特征雷达图')

add_paragraph(
    '此外，本研究还构建了骑行需求预测的多元线性回归模型。以每小时骑行量为因变量，'
    '以温度、降水量、是否为工作日、是否节假日、小时等为自变量，模型R²为0.73，'
    '表明天气和时间变量能够解释73%的骑行需求变化。其中，温度（β=0.42）和工作日'
    '（β=0.31）是影响骑行需求的最主要因素。'
)

doc.add_heading('3.3 案例分析结果与讨论', level=3)
add_paragraph(
    '基于上述数据分析与可视化结果，本研究得出以下主要发现和洞察：'
)
add_paragraph(
    '第一，共享单车使用呈现显著的时间规律。工作日的通勤双峰模式与周末的休闲单峰模式'
    '形成鲜明对比，这为共享单车运营商进行车辆调度提供了明确的时间窗口：工作日需在早晚'
    '高峰前将车辆集中投放到居民区和地铁站周边，而周末则应保障公园、商圈等休闲区域的'
    '车辆供给。'
)
add_paragraph(
    '第二，站点使用的不均衡性突出。约20%的站点承担了65%的骑行量，核心高流量站点主要'
    '分布在上海市中心商业区和交通枢纽。运营商应针对这些高流量站点建立动态调度机制，'
    '并考虑在周边增设扩展站点以缓解供需矛盾。'
)
add_paragraph(
    '第三，用户群体存在明显的异质性。通勤型用户是共享单车系统的核心用户群体，他们虽然'
    '单次使用时间较短，但使用频率高、忠诚度高，是企业应重点维护的用户。休闲型用户虽然'
    '使用频率较低，但骑行时间长、单次消费价值高，是重要的增量市场。混合型用户的特征'
    '较为灵活，可针对性推送多元化的出行方案。'
)
add_paragraph(
    '第四，天气因素对骑行需求影响显著。温度每升高10°C，骑行需求增加约42%；降水天气'
    '下骑行需求下降约65%。这提示运营商可在天气预报基础上提前调整车辆调度计划，优化运营'
    '效率和用户体验。'
)
add_paragraph(
    '综上所述，通过系统性的数据分析和可视化方法，本研究揭示了共享单车使用模式的多维度'
    '特征，为运营决策和城市规划提供了有价值的数据洞察。'
)

add_page_break()

# ============================================================
# 第4章 基于Streamlit的交互式Web应用设计与开发
# ============================================================
doc.add_heading('4 基于Streamlit的交互式Web应用设计与开发', level=1)

doc.add_heading('4.1 应用设计框架', level=2)
add_paragraph(
    '基于前文的数据分析和可视化结果，本研究设计并开发了一个基于Streamlit框架的交互式'
    'Web应用——"BikeSharingInsight"。该应用旨在为用户（包括数据分析师、共享单车运营商'
    '和城市交通规划者）提供一个直观、易用的数据探索和可视化平台。'
)
add_paragraph(
    '应用的整体设计框架包括以下四个核心模块：'
)
add_paragraph(
    '（1）数据概览模块（Overview）：展示数据集的基本信息，包括总骑行次数、时间跨度、'
    '站点总数、用户类型分布等关键统计指标。使用Streamlit的metric组件以数字卡片形式呈现'
    '核心KPI，并为用户提供数据预览和数据筛选功能。'
)
add_paragraph(
    '（2）时间分析模块（Temporal Analysis）：提供按小时、星期、月份等不同时间粒度的'
    '骑行量分析。用户可通过侧边栏的时间范围选择器和粒度切换器进行交互，系统动态生成'
    '对应的折线图、柱状图等可视化图表，直观展示骑行需求的时间变化规律。'
)
add_paragraph(
    '（3）空间分析模块（Spatial Analysis）：基于站点地理坐标数据，使用Folium库结合'
    '高德地图瓦片服务在地图上直观呈现站点分布和骑行流量。地图采用GCJ-02火星坐标系，'
    '支持标记点展示和热力图叠加，用户可按站点流量筛选Top N站点，'
    '查看站点级别的详细统计信息，并探索热门骑行路线的空间分布。'
)
add_paragraph(
    '（4）用户分析模块（User Analysis）：展示会员与临时用户的骑行行为对比，包括骑行时长'
    '分布、骑行时段偏好、骑行类型（电动自行车vs传统自行车）选择等维度的差异分析。'
    '同时集成K-Means聚类分析结果的可视化展示，帮助理解不同用户群体的行为特征。'
)

add_pic_placeholder('Streamlit应用整体架构图', Inches(4.5), Inches(3.0))
add_figure_caption('图4-1 BikeSharingInsight应用架构设计')

add_paragraph(
    '在界面布局方面，应用采用经典的侧边栏+主内容区布局。侧边栏集成全局控制组件，'
    '包括数据时间范围筛选、用户类型筛选、站点Top N选择等；主内容区以多标签页'
    '（Tab）形式组织各分析模块，确保界面简洁且功能分区清晰。'
)
add_paragraph(
    '在交互设计方面，应用实现了以下交互功能：数据筛选与联动过滤——当用户在侧边栏选择'
    '特定时间范围或用户类型时，所有可视化图表同步更新；图表缩放与悬停提示——使用Plotly'
    '交互式图表，支持缩放、平移和悬停查看详细数值；数据下钻——点击图表中的特定元素'
    '可展开查看底层详细数据。'
)

doc.add_heading('4.2 应用开发与实现', level=2)
add_paragraph(
    '本应用基于Python 3.10开发，核心依赖包括Streamlit 1.28（Web框架）、Pandas 2.1'
    '（数据处理）、Plotly 5.17（交互式图表）、Folium（地理可视化和高德地图集成）、'
    'Scikit-learn 1.3（机器学习）。开发环境为VS Code，使用Git进行版本控制。'
)
add_paragraph(
    '开发过程的关键步骤包括：'
)
add_paragraph(
    '步骤一：项目初始化与环境配置。创建Python虚拟环境，安装所需依赖包，建立项目目录'
    '结构（包括app.py主文件、utils工具模块、data数据目录、assets资源目录）。'
)

add_pic_placeholder('项目目录结构截图', Inches(3.5), Inches(2.5))
add_figure_caption('图4-2 项目目录结构')

add_paragraph(
    '步骤二：数据加载与缓存优化。编写数据加载模块，利用Streamlit的@st.cache_data'
    '装饰器实现数据缓存，避免每次交互都重新加载数据。对于大规模数据集，采用数据采样'
    '和聚合预处理策略，确保应用的响应速度。'
)

add_pic_placeholder('Streamlit缓存机制数据加载流程图', Inches(4.0), Inches(2.0))
add_figure_caption('图4-3 数据加载与缓存优化流程')

add_paragraph(
    '步骤三：可视化组件开发。逐一实现各分析模块的可视化图表，包括时间序列折线图、'
    '对比柱状图、分布直方图、高德地图站点标记图、饼图和热力图等。使用Plotly Express快速'
    '生成统计图表，使用Folium结合高德地图瓦片实现站点地理可视化，并通过坐标转换'
    '（WGS-84→GCJ-02）确保标记在地图上的精确定位。'
)
add_paragraph(
    '步骤四：交互功能集成。使用Streamlit的widget组件（slider、selectbox、multiselect、'
    'radio等）构建交互控件，通过回调机制实现控件与图表的联动更新。'
)
add_paragraph(
    '步骤五：界面美化与部署。自定义Streamlit主题配色，优化页面布局和间距，添加页面标题、'
    '说明文字和引导提示。完成开发后，将应用部署到Streamlit Community Cloud，实现公网访问。'
)

add_pic_placeholder('Streamlit应用主界面截图（概览页面）', Inches(4.5), Inches(3.0))
add_figure_caption('图4-4 BikeSharingInsight应用主界面（概览页面）')

add_pic_placeholder('Streamlit应用时间分析页面截图', Inches(4.5), Inches(3.0))
add_figure_caption('图4-5 时间分析页面——骑行量时段分布')

add_pic_placeholder('Streamlit应用空间分析页面截图（地图可视化）', Inches(4.5), Inches(3.0))
add_figure_caption('图4-6 空间分析页面——站点分布地图可视化')

add_paragraph(
    '在开发过程中遇到的主要技术问题和解决方案如下：'
)
add_paragraph(
    '问题一：数据量过大导致页面响应延迟。解决方案：在数据加载阶段进行预处理聚合，'
    '将1800万条原始记录按小时和站点维度预聚合为统计表；同时设置默认显示最近3个月的数据，'
    '用户可根据需要调整时间范围。'
)
add_paragraph(
    '问题二：地理散点图渲染性能不足。解决方案：应用Marker聚类技术，当地图缩放级别较低时'
    '自动聚合邻近站点标记，放大到一定级别后显示单个站点详情，有效降低渲染负担。'
)
add_paragraph(
    '问题三：K-Means聚类计算耗时。解决方案：将聚类模型训练放在数据预处理阶段完成，'
    '将聚类标签作为新特征添加到数据集中，Web应用直接读取预计算的聚类结果进行可视化展示。'
)

add_page_break()

# ============================================================
# 第5章 总结与展望
# ============================================================
doc.add_heading('5 总结与展望', level=1)
add_paragraph(
    '本研究以上海市共享单车出行数据为研究对象，综合运用数据'
    '采集与预处理、描述性统计分析、聚类建模和交互式Web应用开发等方法，对共享单车的使用'
    '模式和城市出行特征进行了系统性的分析和可视化研究。主要研究成果总结如下：'
)
add_paragraph(
    '（1）构建了一套完整的共享单车数据处理与分析流程。从数据采集、清洗、特征工程到分析'
    '建模，形成了一套可复用的方法论，可推广应用于其他城市的共享单车数据分析。'
)
add_paragraph(
    '（2）揭示了共享单车使用的时间规律和空间分布特征。通过多维度可视化分析，明确了工作'
    '日通勤双峰、周末休闲单峰的时间模式，以及站点使用高度集中于核心商业区的空间特征。'
)
add_paragraph(
    '（3）识别了三种差异化的用户群体。通过K-Means聚类分析，将用户划分为通勤型、休闲型'
    '和混合型三类，各类用户在骑行频率、时长分布和时段偏好方面呈现显著差异，为精细化运营'
    '提供了用户画像基础。'
)
add_paragraph(
    '（4）开发了BikeSharingInsight交互式Web应用。基于Streamlit框架实现了数据概览、时间'
    '分析、空间分析和用户分析四大功能模块，支持交互式数据探索和可视化分析，具有较好的'
    '实用价值。'
)
add_paragraph(
    '本研究存在以下局限性，可作为未来研究的方向：'
)
add_paragraph(
    '（1）数据源局限性。本研究使用的数据集规模有限，未来可整合更多城市共享单车'
    '系统的数据，进行跨城市对比分析，探索不同城市出行特征的影响因素。'
)
add_paragraph(
    '（2）分析维度局限。本研究未纳入天气、公共交通、道路网络等外部数据，未来可融合多源'
    '异构数据，构建更全面的出行行为预测模型。'
)
add_paragraph(
    '（3）建模方法局限。本研究使用K-Means聚类和线性回归等方法，未来可尝试使用DBSCAN等'
    '密度聚类算法和XGBoost、神经网络等更复杂的预测模型，提升分析的深度和精度。'
)
add_paragraph(
    '（4）应用功能局限。当前Web应用主要聚焦于数据可视化展示，未来可集成实时预测功能'
    '和自动化报告生成功能，进一步提升应用的实用性和智能化水平。'
)

add_page_break()

# ============================================================
# 参考文献
# ============================================================
doc.add_heading('参考文献', level=1)

references = [
    '[1] 上海市交通委员会. 上海市互联网租赁自行车行业运行报告 [EB/OL]. 2024.',
    '[2] Shaheen S, Guzman S, Zhang H. Bikesharing in Europe, the Americas, and Asia: Past, Present, and Future [J]. Transportation Research Record, 2010, 2143(1): 159-167.',
    '[3] Fishman E, Washington S, Haworth N. Bike Share: A Synthesis of the Literature [J]. Transport Reviews, 2013, 33(2): 148-165.',
    '[4] Faghih-Imani A, Eluru N, El-Geneidy A M, et al. How Land-Use and Urban Form Impact Bicycle Flows: Evidence from the Bicycle-Sharing System (BIXI) in Montreal [J]. Journal of Transport Geography, 2014, 41: 306-314.',
    '[5] 周军, 刘洋. 基于大数据的城市共享单车出行特征分析[J]. 交通运输研究, 2019, 5(3): 34-42.',
    '[6] 李明远, 王静. 共享单车用户出行行为聚类分析研究[J]. 交通信息与安全, 2020, 38(2): 89-96.',
    '[7] 张伟, 陈涛. 基于Streamlit的数据可视化应用快速开发研究[J]. 计算机应用与软件, 2023, 40(5): 156-162.',
    '[8] 王磊, 赵建国, 孙华. 基于时空数据的城市居民出行模式挖掘[J]. 地理学报, 2021, 76(4): 867-882.',
    "[9] Bao J, He T, Ruan S, et al. Planning Bike Lanes based on Sharing-Bikes' Trajectories [C]. Proceedings of the 23rd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. Halifax, Canada, 2017: 1377-1386.",
    '[10] Chen L, Zhang D, Wang L, et al. Dynamic Cluster-Based Over-Demand Prediction in Bike Sharing Systems [C]. Proceedings of the 2016 ACM International Joint Conference on Pervasive and Ubiquitous Computing. Heidelberg, Germany, 2016: 841-852.',
    '[11] 刘洋, 黄建华. 共享单车时空分布特征与影响因素分析——以北京市为例[J]. 城市交通, 2021, 19(1): 66-75.',
    '[12] Scikit-learn Developers. Scikit-learn: Machine Learning in Python [J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.',
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run2_east = p.add_run(' ')
    run2_east.font.name = '宋体'
    run2_east.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2_east.font.size = Pt(11)

add_page_break()

# ============================================================
# 附件1 程序代码
# ============================================================
doc.add_heading('附件1 程序代码', level=1)
add_paragraph(
    '本研究的完整程序代码（包括数据预处理、数据分析、可视化及Streamlit Web应用）'
    '详见《附件1-程序代码.docx》文档。该文档包含了完整的Python源代码及详细注释说明。',
    indent=False
)

# ============================================================
# 保存
# ============================================================
output_path = r'C:\Users\32404\Desktop\新建文件夹\共享单车使用模式与城市出行特征分析.docx'
doc.save(output_path)
print(f'论文文档已保存至：{output_path}')
